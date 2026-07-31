import os
import pypdf
import docx

def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
    except Exception as e:
        text = f"Error reading PDF {os.path.basename(file_path)}: {str(e)}"
    return text

def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                tables_text.append(" | ".join(row_text))
        text = "\n".join(paragraphs + tables_text)
    except Exception as e:
        text = f"Error reading DOCX {os.path.basename(file_path)}: {str(e)}"
    return text

def parse_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"Error reading TXT {os.path.basename(file_path)}: {str(e)}"

def extract_text_from_file(file_path: str) -> str:
    """Detect file extension and extract text accordingly."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in [".txt", ".md", ".json"]:
        return parse_txt(file_path)
    else:
        # Fallback to text reading if unknown
        return parse_txt(file_path)
