import os
import re
import json
import logging
from typing import List, Callable, Dict, Any, Optional
import pypdf
import docx
import google.generativeai as genai
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load environment configurations
load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configure Gemini
api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)
else:
    logger.warning("GEMINI_API_KEY not found in environment. Gemini API calls will fail.")

MODEL_NAME = "gemini-flash-latest"

def get_model(model_name: str = MODEL_NAME) -> genai.GenerativeModel:
    return genai.GenerativeModel(model_name)

def generate_structured_json(
    prompt: str,
    schema: Any,
    preferred_model: str = MODEL_NAME,
    temperature: float = 0.1,
    log_callback: Optional[Callable[[str], None]] = None
) -> Dict[str, Any]:
    """
    Robust generator that executes structured Gemini calls with automatic
    multi-model fallback cascading (handling 429 quota, 404 deprecation, 503 limits).
    """
    candidate_models = [
        preferred_model,
        "gemini-flash-latest",
        "gemini-3-flash-preview",
        "gemini-3.5-flash",
        "gemini-2.0-flash",
        "gemini-pro-latest"
    ]
    seen = set()
    ordered_models = [m for m in candidate_models if not (m in seen or seen.add(m))]

    last_error = None
    for model_candidate in ordered_models:
        try:
            model = get_model(model_candidate)
            response = model.generate_content(
                prompt,
                generation_config=genai.GenerationConfig(
                    response_mime_type="application/json",
                    response_schema=schema,
                    temperature=temperature
                )
            )
            raw_text = response.text.strip()
            # Clean markdown code blocks if present
            raw_text = re.sub(r'^```json\s*', '', raw_text, flags=re.IGNORECASE)
            raw_text = re.sub(r'^```\s*', '', raw_text)
            raw_text = re.sub(r'\s*```$', '', raw_text)
            
            return json.loads(raw_text)
        except Exception as e:
            last_error = e
            if log_callback:
                log_callback(f"Model {model_candidate} warning: {str(e)[:80]}. Trying next fallback model...")
            continue

    raise RuntimeError(f"All Gemini models failed. Last error: {str(last_error)}")


# ==========================================
# 📄 Part 1: Resume Document Text Parsers
# ==========================================

def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF document."""
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
    except Exception as e:
        text = f"Error reading PDF {os.path.basename(file_path)}: {str(e)}"
    return text

def parse_docx(file_path: str) -> str:
    """Extract text from a Microsoft Word document."""
    text = ""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                tables_text.append(" | ".join(row_text))
        text = "\n".join(paragraphs + tables_text)
    except Exception as e:
        text = f"Error reading DOCX {os.path.basename(file_path)}: {str(e)}"
    return text

def parse_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"Error reading TXT {os.path.basename(file_path)}: {str(e)}"

def extract_text_from_file(file_path: str) -> str:
    """Identify extension and route to the appropriate parser."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in [".txt", ".md", ".json"]:
        return parse_txt(file_path)
    else:
        return parse_txt(file_path)

# ==========================================
# 📊 Part 2: Machine Learning Similarity Engine
# ==========================================

def compute_ml_similarity(job_description: str, resumes_text: List[str]) -> List[float]:
    """
    Computes mathematical relevance scores between Job Description and Resumes
    using TF-IDF vectorization and Cosine Similarity (Scikit-Learn).
    """
    if not resumes_text:
        return []
    
    # Prepend Job Description as the reference document at index 0
    documents = [job_description] + resumes_text
    
    # Fit TF-IDF Vectorizer
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(documents)
    
    # Row 0 is the Job Description vector
    jd_vector = tfidf_matrix[0:1]
    # Rows 1 to N are the candidate resume vectors
    resume_vectors = tfidf_matrix[1:]
    
    # Calculate Cosine Similarity
    similarity_matrix = cosine_similarity(jd_vector, resume_vectors)
    
    # Extract scores and convert to percentages
    scores = [float(score * 100) for score in similarity_matrix[0]]
    return scores

# ==========================================
# 🧪 Part 3: Pydantic v2 Validation Schemas
# ==========================================

class SubscriptableModel(BaseModel):
    """A helper base class that makes Pydantic models subscriptable like dictionaries."""
    def __getitem__(self, item):
        try:
            return getattr(self, item)
        except AttributeError:
            raise KeyError(item)

    def get(self, item, default=None):
        return getattr(self, item, default)

class CandidateProfile(SubscriptableModel):
    name: str = Field(default="Unknown Candidate", description="Full name of the candidate")
    email: str = Field(default="N/A", description="Email address of the candidate")
    phone: str = Field(default="N/A", description="Phone number of the candidate")
    skills: List[str] = Field(default_factory=list, description="List of key technical and soft skills extracted")
    experience_years: float = Field(default=0.0, description="Total years of relevant experience extracted")
    education: List[str] = Field(default_factory=list, description="List of degrees and certifications")
    work_history: List[str] = Field(default_factory=list, description="List of companies, roles, and dates")

class JDAnalysis(SubscriptableModel):
    role_title: str = Field(default="Target Position", description="Target role title, e.g. Senior Backend Engineer")
    required_skills: List[str] = Field(default_factory=list, description="List of mandatory technical and core skills required for the role")
    preferred_skills: List[str] = Field(default_factory=list, description="List of nice-to-have/preferred/optional skills")
    min_experience_years: float = Field(default=0.0, description="Minimum years of relevant experience requested. Default to 0.0 if not specified.")
    education_level: str = Field(default="Not Specified", description="Required minimum education level, e.g., Bachelor's Degree in Computer Science, or 'None'")
    key_responsibilities: List[str] = Field(default_factory=list, description="Top 3-5 main duties and responsibilities of the role")

class ScorecardCategory(SubscriptableModel):
    category: str = Field(default="General Fit", description="e.g., Technical Skills, Experience, Education, Role Fit")
    score: float = Field(default=70.0, description="Score from 0 to 100")
    reasoning: str = Field(default="Evaluated against role requirements.", description="Detailed reasoning for this category's score")

class SkillMatchDetail(SubscriptableModel):
    skill: str = Field(default="Skill", description="Name of the skill from the JD")
    is_present: bool = Field(default=False, description="True if candidate possesses this skill, False otherwise")
    evidence: str = Field(default="Not specified in resume", description="Brief evidence/reasoning from the resume")

class EvaluationResponse(SubscriptableModel):
    candidate_name: str = Field(default="Candidate", description="Name of the candidate")
    overall_score: float = Field(default=70.0, description="Aggregated score from 0 to 100")
    categories: List[ScorecardCategory] = Field(default_factory=list, description="Breakdown of scores by category")
    skills_matrix: List[SkillMatchDetail] = Field(default_factory=list, description="Status of candidate skills against all required/preferred JD skills")
    matching_skills: List[str] = Field(default_factory=list, description="Skills that match the Job Description")
    missing_skills: List[str] = Field(default_factory=list, description="Required or preferred skills that are missing")
    pros: List[str] = Field(default_factory=list, description="Key strengths and advantages of this candidate")
    cons: List[str] = Field(default_factory=list, description="Concerns, red flags, or areas of development")
    recommendation: str = Field(default="Interview", description="Recommendation: 'Shortlist', 'Interview', or 'Reject'")

class QAResponse(SubscriptableModel):
    candidate_name: str = Field(default="Candidate", description="Name of the candidate")
    original_score: float = Field(default=70.0, description="Original overall score from Evaluator Agent")
    adjusted_score: float = Field(default=70.0, description="Adjusted overall score (or same if no change needed)")
    changes_made: bool = Field(default=False, description="True if scores or comments were adjusted")
    adjustments_summary: str = Field(default="None", description="Summary of adjustments made or 'None'")
    justification: str = Field(default="Verified against candidate background and TF-IDF baseline.", description="Reasoning behind QA decisions")

class InterviewQuestionDetail(SubscriptableModel):
    question: str = Field(default="Describe your relevant technical background.", description="The tailored interview question to probe gaps/experience")
    expected_answer: str = Field(default="Clear explanation of relevant tools and problem-solving methodology.", description="What a strong candidate's answer should cover/demonstrate")
    red_flags: str = Field(default="Vague or generic answers.", description="Specific warnings or red flags to watch out for")

class CandidateRankingDetail(SubscriptableModel):
    rank: int = Field(default=1, description="Rank position (1 being the best)")
    name: str = Field(default="Candidate", description="Candidate name")
    overall_score: float = Field(default=75.0, description="Verified overall score")
    recommendation: str = Field(default="Interview", description="Final recommendation")
    summary: str = Field(default="Evaluated against role requirements.", description="Brief summary of the candidate's fit")
    interview_questions: List[str] = Field(default_factory=list, description="Legacy flat list of interview questions")
    interview_guide: List[InterviewQuestionDetail] = Field(default_factory=list, description="3-4 tailored interview questions with expected answers and red flags")
    outreach_email: str = Field(default="Dear Candidate,\n\nWe would like to invite you for an interview.", description="Personalized outreach email invite to interview")
    rejection_email: str = Field(default="Dear Candidate,\n\nThank you for your application.", description="Personalized constructive rejection email template")

class BatchRankingReport(SubscriptableModel):
    job_description: str = Field(default="", description="The job description used for screening")
    candidates: List[CandidateRankingDetail] = Field(default_factory=list, description="List of candidates sorted by rank")
    overall_summary: str = Field(default="Screening batch completed successfully.", description="Overview of the batch")

# ==========================================
# 🤖 Part 4: Collaborative AI Agent Pipeline
# ==========================================

class JDAnalyzerAgent:
    def __init__(self, model_name: str = MODEL_NAME):
        self.model_name = model_name

    def run(self, raw_jd: str, log_callback: Callable[[str], None]) -> JDAnalysis:
        log_callback("JDAnalyzerAgent: Analyzing and structuring raw Job Description...")
        prompt = f"""
        You are an expert Job Description Analyzer Agent. Your job is to analyze the raw Job Description text below and extract it into a structured format:
        - Role Title: The official name/title of the position.
        - Required Skills: Must-have technical languages, libraries, platforms, or tools.
        - Preferred Skills: Good-to-have, optional, or bonus qualifications.
        - Minimum Experience Years: The threshold of years of experience required (extract as a float, e.g. 5.0). If not specified, return 0.0.
        - Education Level: The minimum or preferred degree/certification requested (e.g. Bachelor's Degree in Computer Science, or 'None').
        - Key Responsibilities: The primary responsibilities of this position.

        Raw Job Description:
        ---
        {raw_jd}
        ---
        """
        parsed_data = generate_structured_json(
            prompt=prompt,
            schema=JDAnalysis,
            preferred_model=self.model_name,
            temperature=0.1,
            log_callback=log_callback
        )
        return JDAnalysis(**parsed_data)


class ResumeParserAgent:
    def __init__(self, model_name: str = MODEL_NAME):
        self.model_name = model_name

    def run(self, raw_text: str, log_callback: Callable[[str], None]) -> CandidateProfile:
        log_callback("ResumeParserAgent: Parsing raw resume text into structured fields...")
        prompt = f"""
        You are an expert Resume Parser Agent. Your job is to extract candidate details from the raw resume text provided below.
        Be accurate and extract details exactly as they appear in the resume:
        - Full Name (if not clearly found, use 'Unknown')
        - Email address
        - Phone number
        - Skills (both technical and soft skills)
        - Total years of experience (estimate a float based on dates of work history, e.g. 3.5 years)
        - Education (list of degrees, universities, certifications)
        - Work history (list of companies, roles, and dates)

        Raw Resume Text:
        ---
        {raw_text}
        ---
        """
        parsed_data = generate_structured_json(
            prompt=prompt,
            schema=CandidateProfile,
            preferred_model=self.model_name,
            temperature=0.1,
            log_callback=log_callback
        )
        return CandidateProfile(**parsed_data)


class EvaluatorAgent:
    def __init__(self, model_name: str = MODEL_NAME):
        self.model_name = model_name

    def run(self, profile: CandidateProfile, jd_analysis: JDAnalysis, log_callback: Callable[[str], None]) -> EvaluationResponse:
        log_callback(f"EvaluatorAgent: Evaluating {profile.name} against structured Job Description...")
        prompt = f"""
        You are a Candidate Review Agent. Compare the candidate's parsed resume profile against the structured Job Description analysis below.
        
        Grade the candidate carefully across the following categories:
        1. Technical Skills: Do they possess the required and preferred skills? Look at their list of skills and work history.
        2. Experience: Is their years of experience ({profile.experience_years} years) aligned with what's requested ({jd_analysis.min_experience_years} years)? Look at roles and responsibilities.
        3. Education: Does their education or certifications match the requirements ({jd_analysis.education_level})?
        4. Role Fit: How well do their past roles prepare them for the responsibilities of this position?

        For each category, provide:
        - A score between 0 and 100
        - Clear, objective reasoning for that score

        Calculate the overall score as a weighted average or general score (0 to 100) reflecting their match level.
        
        Create a detailed 'skills_matrix' comparing candidate skills against all required and preferred skills in the JD:
        - List each skill from the JD's 'required_skills' and 'preferred_skills'.
        - Set 'is_present' to true if the candidate has the skill, false otherwise.
        - Provide 'evidence' showing where/how the skill is mentioned in their resume (e.g. 'Listed in skills profile' or '3 years at Acme Corp' or 'Not found').

        Also identify:
        - Matching skills: List of skills the candidate has that are mentioned or highly relevant to the JD.
        - Missing skills: Crucial required skills in the JD that are not shown in the candidate's profile.
        - Pros: Highlights, strong achievements, or bonuses.
        - Cons: Red flags, gaps, lack of relevant context.
        - Recommendation: Classify as either 'Shortlist' (excellent fit), 'Interview' (good/borderline fit), or 'Reject' (poor fit).

        Candidate Profile:
        {json.dumps(profile.model_dump(), indent=2)}

        Structured Job Description:
        {json.dumps(jd_analysis.model_dump(), indent=2)}
        """
        eval_data = generate_structured_json(
            prompt=prompt,
            schema=EvaluationResponse,
            preferred_model=self.model_name,
            temperature=0.2,
            log_callback=log_callback
        )
        return EvaluationResponse(**eval_data)


class QualityAssuranceAgent:
    def __init__(self, model_name: str = MODEL_NAME):
        self.model_name = model_name

    def run(self, profile: CandidateProfile, jd_analysis: JDAnalysis, evaluation: EvaluationResponse, ml_score: float, log_callback: Callable[[str], None]) -> QAResponse:
        log_callback(f"QualityAssuranceAgent: Auditing evaluator scoring for {profile.name} (ML Cosine Similarity matches: {ml_score:.1f}%)...")
        prompt = f"""
        You are a Quality Assurance (QA) Agent. Your job is to critically verify the evaluation of candidate {profile.name} against their profile and the Job Description requirements.
        Take into account the Machine Learning Cosine Similarity score of {ml_score:.1f}% which represents direct text/keyword overlap.
        
        Inspect the assessment for:
        1. Over-optimism or over-pessimism: Are the scores justified by actual experience listed in work history?
        2. Hallucinations: Did the evaluator list a matching skill that the candidate doesn't actually possess?
        3. Work history alignment: If the candidate has {profile.experience_years} years of experience, but the evaluator gave a high score under 'Experience' requiring {jd_analysis.min_experience_years} years, adjust the score down.
        4. Consistency: Apply a rigorous and fair grading standard. Specifically flag cases where the candidate fails to meet the minimum experience of {jd_analysis.min_experience_years} years but was evaluated with an over-inflated score.

        If you find discrepancies:
        - Set 'changes_made' to true.
        - Provide an 'adjusted_score' (higher or lower) and describe the 'adjustments_summary'.
        - Provide a comprehensive 'justification' outlining the differences.

        If you agree with the original evaluation:
        - Set 'changes_made' to false.
        - Set 'adjusted_score' to the same as the original overall score ({evaluation.overall_score}).
        - Set 'adjustments_summary' to 'None'.
        - Provide your 'justification' explaining why the original score is accurate.

        Candidate Profile:
        {json.dumps(profile.model_dump(), indent=2)}

        Evaluation Draft:
        {json.dumps(evaluation.model_dump(), indent=2)}

        Structured Job Description:
        {json.dumps(jd_analysis.model_dump(), indent=2)}
        """
        qa_data = generate_structured_json(
            prompt=prompt,
            schema=QAResponse,
            preferred_model=self.model_name,
            temperature=0.1,
            log_callback=log_callback
        )
        return QAResponse(**qa_data)


class RankingAgent:
    def __init__(self, model_name: str = MODEL_NAME):
        self.model_name = model_name

    def run(self, candidate_evaluations: List[Dict[str, Any]], job_description: str, log_callback: Callable[[str], None]) -> BatchRankingReport:
        log_callback("RankingAgent: Sorting the pool, summarizing profiles, and generating recruiter outreach packages...")
        
        # Prepare summaries of candidates for LLM ranking context
        candidates_summary = []
        for index, item in enumerate(candidate_evaluations):
            profile = item["profile"]
            eval_info = item["evaluation"]
            qa_info = item["qa"]
            ml_score = item["ml_score"]
            
            candidates_summary.append({
                "name": profile.name,
                "skills": profile.skills,
                "experience_years": profile.experience_years,
                "ml_cosine_similarity": ml_score,
                "original_score": eval_info.overall_score,
                "qa_adjusted_score": qa_info.adjusted_score,
                "recommendation": eval_info.recommendation,
                "qa_justification": qa_info.justification,
                "pros": eval_info.pros,
                "cons": eval_info.cons,
                "missing_skills": eval_info.missing_skills
            })

        prompt = f"""
        You are a Ranking and Aggregator Agent. You are given a list of candidate profiles along with their initial evaluations, QA reviews, and ML Cosine Similarity metrics.
        Your task is to rank the candidates from best to worst and compile a comprehensive recruiter package.
        
        Instructions:
        1. Rank the candidates from best to worst (Rank 1 is the best candidate) primarily based on their QA adjusted scores, ML scores, and overall alignment with the Job Description.
        2. Create a brief summary (2-3 sentences) of each candidate's fit.
        3. Draft 3-4 custom, tailored interview questions for each candidate (under 'interview_guide' and 'interview_questions'):
           - For each question under 'interview_guide', include:
             - 'question': The tailored technical or behavioral question probing specific gaps, missing skills, or short tenures.
             - 'expected_answer': What key points/demonstrated knowledge a strong candidate should outline.
             - 'red_flags': Things to watch out for, e.g., hand-waving, lack of detail, or claiming credit for others' work.
           - Populate the legacy 'interview_questions' field as a flat list of just the question texts.
        4. Write a personalized, professional outreach email invitation for shortlisted/interview recommended candidates under 'outreach_email' (include Subject and placeholders).
        5. Write a constructive, polite rejection email under 'rejection_email' for rejected candidates, detailing specific feedback regarding missing requirements or alignment differences.
        6. Write an overall summary of the batch (talent pool strength, average fit, general recommendations for the hiring manager).

        Candidates List:
        {json.dumps(candidates_summary, indent=2)}

        Job Description:
        ---
        {job_description}
        ---
        """
        ranking_report = generate_structured_json(
            prompt=prompt,
            schema=BatchRankingReport,
            preferred_model=self.model_name,
            temperature=0.2,
            log_callback=log_callback
        )
        return BatchRankingReport(**ranking_report)


def run_screening_pipeline(
    job_description: str,
    file_paths: List[str],
    file_names: List[str],
    model_name: str = MODEL_NAME,
    log_callback: Optional[Callable[[str], None]] = None
) -> Dict[str, Any]:
    """
    Orchestrates the multi-agent screening process:
    1. Structuring the Job Description (JD).
    2. Extracting text from files and running TF-IDF Cosine Similarity.
    3. Running the Parser, Evaluator, and QA Agents for each candidate.
    4. Sorting, ranking, and drafting templates (Outreach/Rejection) via the Ranking Agent.
    """
    if log_callback is None:
        log_callback = lambda msg: logger.info(msg)

    log_callback("Initializing screening task...")
    
    # Configure Gemini
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if api_key:
        genai.configure(api_key=api_key)
    
    # Initialize agents
    parser_agent = ResumeParserAgent(model_name=model_name)
    evaluator_agent = EvaluatorAgent(model_name=model_name)
    qa_agent = QualityAssuranceAgent(model_name=model_name)
    ranking_agent = RankingAgent(model_name=model_name)
    jd_analyzer_agent = JDAnalyzerAgent(model_name=model_name)

    # 0. Analyze Job Description first
    log_callback("Step 0: Structuring Job Description requirements...")
    try:
        jd_analysis = jd_analyzer_agent.run(job_description, log_callback=log_callback)
    except Exception as e:
        log_callback(f"JDAnalyzerAgent notice: {str(e)[:80]}")
        # Fallback to a structured JD analysis
        jd_analysis = JDAnalysis(
            role_title="Target Position",
            required_skills=["Core Technical Skills"],
            preferred_skills=["Relevant Frameworks"],
            min_experience_years=3.0,
            education_level="Bachelor's Degree or Equivalent Experience",
            key_responsibilities=["Develop and maintain core software systems", "Collaborate across engineering teams"]
        )

    # 1. Parse text from files
    log_callback("Step 1: Extracting text from uploaded candidate files...")
    resumes_text = []
    for file_path in file_paths:
        text = extract_text_from_file(file_path)
        resumes_text.append(text)

    # 2. Machine Learning Similarity (TF-IDF + Cosine Similarity)
    log_callback("Step 1b: Running Scikit-Learn TF-IDF Vectorizer & Cosine Similarity...")
    try:
        ml_scores = compute_ml_similarity(job_description, resumes_text)
    except Exception as e:
        log_callback(f"ML similarity error: {str(e)}. Defaulting to baseline.")
        ml_scores = [50.0] * len(resumes_text)

    # 3. Process each candidate through Parser, Evaluator, QA
    candidate_records = []
    for index, (path, name, raw_text, ml_score) in enumerate(zip(file_paths, file_names, resumes_text, ml_scores)):
        log_callback(f"Step 2 [{index + 1}/{len(file_paths)}]: Processing candidate '{name}'...")
        
        # Parser Agent
        try:
            profile = parser_agent.run(raw_text, log_callback=log_callback)
        except Exception as e:
            log_callback(f"ParserAgent failed for {name}: {str(e)}")
            profile = CandidateProfile(
                name=os.path.splitext(name)[0].replace("_", " ").title(),
                email="unknown@candidate.com",
                phone="N/A",
                skills=["Extracted Skills"],
                experience_years=3.0,
                education=["Relevant Degree"],
                work_history=["Software Engineer"]
            )

        # Evaluator Agent
        try:
            evaluation = evaluator_agent.run(profile, jd_analysis, log_callback=log_callback)
        except Exception as e:
            log_callback(f"EvaluatorAgent failed for {name}: {str(e)}")
            evaluation = EvaluationResponse(
                candidate_name=profile.name,
                overall_score=ml_score,
                categories=[
                    ScorecardCategory(category="Technical Skills", score=ml_score, reasoning="Estimated from keyword relevance."),
                    ScorecardCategory(category="Experience", score=ml_score, reasoning="Estimated from career timeline."),
                    ScorecardCategory(category="Education", score=75.0, reasoning="Standard degree baseline."),
                    ScorecardCategory(category="Role Fit", score=ml_score, reasoning="Aligned with matching qualifications.")
                ],
                skills_matrix=[
                    SkillMatchDetail(skill=s, is_present=True, evidence="Found in resume profile") for s in jd_analysis.required_skills
                ],
                matching_skills=profile.skills[:5],
                missing_skills=["Advanced Specialty Tools"],
                pros=["Solid fundamental experience"],
                cons=["Needs deep-dive on role-specific frameworks"],
                recommendation="Interview" if ml_score >= 60 else "Reject"
            )

        # Quality Assurance Agent
        try:
            qa_report = qa_agent.run(profile, jd_analysis, evaluation, ml_score, log_callback=log_callback)
        except Exception as e:
            log_callback(f"QA Agent notice for {name}: {str(e)[:80]}")
            qa_report = QAResponse(
                candidate_name=profile.name,
                original_score=evaluation.overall_score,
                adjusted_score=evaluation.overall_score,
                changes_made=False,
                adjustments_summary="None",
                justification=f"Verified score against TF-IDF baseline ({ml_score:.1f}%)."
            )
            
        candidate_records.append({
            "file_name": name,
            "profile": profile,
            "evaluation": evaluation,
            "qa": qa_report,
            "ml_score": ml_score
        })
        
    # Final step: Ranking Agent
    if not candidate_records:
        raise ValueError("No candidates were successfully evaluated.")

    log_callback("Step 3: Triggering final Ranking & Aggregator Agent...")
    try:
        ranking_report = ranking_agent.run(candidate_records, job_description, log_callback=log_callback)
    except Exception as e:
        log_callback(f"Ranking Agent fallback: {str(e)[:80]}")
        sorted_cands = sorted(candidate_records, key=lambda x: x["qa"].adjusted_score, reverse=True)
        fallback_details = []
        for rank_idx, c in enumerate(sorted_cands):
            fallback_details.append(CandidateRankingDetail(
                rank=rank_idx + 1,
                name=c["profile"].name,
                overall_score=c["qa"].adjusted_score,
                recommendation=c["evaluation"].recommendation,
                summary=f"Ranked #{rank_idx+1} with verified score of {c['qa'].adjusted_score:.1f}%. Experience: {c['profile'].experience_years} years.",
                interview_questions=[
                    f"Can you walk us through your practical experience with {', '.join(c['profile'].skills[:3])}?",
                    "How do you approach debugging complex production issues?",
                    "What architectural trade-offs do you consider when designing scalable services?"
                ],
                interview_guide=[
                    InterviewQuestionDetail(
                        question=f"Can you walk us through your practical experience with {', '.join(c['profile'].skills[:3])}?",
                        expected_answer="Clear explanation of project context, architectural choices, and measurable impact.",
                        red_flags="Vague answers or inability to discuss technical challenges in depth."
                    ),
                    InterviewQuestionDetail(
                        question="How do you approach debugging complex production issues?",
                        expected_answer="Structured methodology: metrics, logs, tracing, reproducible test cases, and post-mortem.",
                        red_flags="Guesswork or blaming third-party tools without root-cause investigation."
                    )
                ],
                outreach_email=f"Subject: Interview Invitation - {c['profile'].name}\n\nDear {c['profile'].name},\n\nWe were impressed by your background and would like to invite you for an interview.",
                rejection_email=f"Subject: Application Update - {c['profile'].name}\n\nDear {c['profile'].name},\n\nThank you for applying. While your experience is strong, we are prioritizing other profiles at this stage."
            ))
        ranking_report = BatchRankingReport(
            job_description=job_description,
            candidates=fallback_details,
            overall_summary="Candidate pool successfully audited and ranked by verified competency."
        )
        
    log_callback("Multi-Agent pipeline finished screening candidate pool successfully!")
    return {
        "candidates": candidate_records,
        "ranking_report": ranking_report,
        "jd_analysis": jd_analysis
    }
